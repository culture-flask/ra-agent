"""多格式文档解析：txt/md 直接读文本，pdf/docx 用轻量库提取。

选型说明：轻量起步用 pypdf / python-docx（纯 Python、零模型下载、可靠）。
生产级解析（版面分析、表格、公式）可换 Marker / Docling——只需在本模块
增加/替换解析函数，上层（入库流水线）零改动。
"""

import io
from pathlib import Path


def parse_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def parse_pdf(content: bytes) -> str:
    return "\n".join(text for _, text in parse_pdf_pages(content))


def parse_docx(content: bytes) -> str:
    return "\n".join(p.text for p in _docx_document(content).paragraphs)


def _docx_document(content: bytes):
    from docx import Document
    return Document(io.BytesIO(content))


# ---------- 逐页解析（chunk 元数据需要源文件名 + 页码） ----------

def parse_txt_pages(content: bytes) -> list[tuple[None, str]]:
    """txt/md：无分页概念，整篇作为一"页"，页码为 None。"""
    return [(None, content.decode("utf-8", errors="replace"))]


def parse_pdf_pages(content: bytes) -> list[tuple[int, str]]:
    """PDF：逐页提取文本，页码从 1 开始（空页保留页码，由调用方过滤）。"""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    return [(i + 1, page.extract_text() or "") for i, page in enumerate(reader.pages)]


def parse_docx_pages(content: bytes) -> list[tuple[None, str]]:
    """docx：Word 分页由渲染决定，无可靠页码，整篇作为一"页"，页码为 None。"""
    doc = _docx_document(content)
    return [(None, "\n".join(p.text for p in doc.paragraphs))]


# 扩展点：新增格式 = 在此登记一个逐页解析函数
PAGE_PARSERS = {
    ".txt": parse_txt_pages,
    ".md": parse_txt_pages,
    ".pdf": parse_pdf_pages,
    ".docx": parse_docx_pages,
}

# 兼容层：全文解析 = 逐页解析的结果拼接
PARSERS = {
    ext: (lambda content, fn=fn: "\n".join(t for _, t in fn(content)))
    for ext, fn in PAGE_PARSERS.items()
}


def parse_file_pages(filename: str, content: bytes) -> list[tuple[int | None, str]]:
    """按扩展名解析文件，返回 [(页码, 文本), ...]。

    - PDF 逐页提取（页码从 1 起，空页被过滤但保留原页码）
    - txt/md/docx 无分页概念 → 单条 (None, 全文)
    不支持的格式抛 ValueError；整篇无文本也抛（扫描件/加密 PDF）。
    """
    ext = Path(filename).suffix.lower()
    parser = PAGE_PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"unsupported file type: {ext} (支持: {', '.join(PAGE_PARSERS)})")
    pages = [(page, text) for page, text in parser(content) if text.strip()]
    if not pages:
        raise ValueError(f"无法从 {filename} 中提取到文本（可能是扫描件/加密 PDF）")
    return pages


def parse_file(filename: str, content: bytes) -> str:
    """按扩展名解析文件，返回提取出的纯文本。不支持的格式抛 ValueError。"""
    return "\n".join(text for _, text in parse_file_pages(filename, content))